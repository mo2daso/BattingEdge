"""
BattingEdge FYP — Project Report Generator
Generates a formatted .docx document of all backend/frontend work and test cases.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'),  'clear')
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, italic=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0xA0, 0xA0)
    return p

def add_table(doc, headers, rows, header_bg='2E4057'):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = tbl.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_bg(hdr[i], header_bg)
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.size = Pt(10)

    for row in rows:
        cells = tbl.add_row().cells
        for i, val in enumerate(row):
            cells[i].paragraphs[0].add_run(str(val)).font.size = Pt(9.5)

    return tbl

def page_break(doc):
    doc.add_page_break()

# ── Build document ────────────────────────────────────────────────────────────

doc = Document()

# Margins
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3)
    sec.right_margin  = Cm(2.5)

# Default font
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(11)

# ─────────────────────────────────────────────────────────────────────────────
# COVER PAGE
# ─────────────────────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('BattingEdge')
r.font.size  = Pt(36)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x00, 0xD4, 0xFF)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run('AI Cricket Batting Analysis System')
r2.font.size = Pt(18)
r2.font.bold = True

doc.add_paragraph()

meta = [
    ('Project Type',  'Final Year Project (FYP)'),
    ('Institution',   'Bahria University, Karachi Campus'),
    ('Department',    'Computer Science'),
    ('Model Version', 'V9.5 Stacking Ensemble — 94.71% Accuracy'),
    ('Backend',       'FastAPI + SQLite + Python 3.11'),
    ('Frontend',      'React 18 + Vite + Tailwind CSS + Framer Motion'),
    ('Defense Date',  'May 7, 2025'),
    ('Report Date',   datetime.date.today().strftime('%B %d, %Y')),
]

for label, val in meta:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_l = p.add_run(f'{label}: ')
    r_l.bold = True
    r_l.font.size = Pt(12)
    r_v = p.add_run(val)
    r_v.font.size = Pt(12)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# TABLE OF CONTENTS (manual)
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, 'Table of Contents', level=1)
toc_items = [
    ('1', 'Project Overview'),
    ('2', 'System Architecture'),
    ('3', 'Backend Development'),
    ('3.1', 'Database Schema'),
    ('3.2', 'Authentication System'),
    ('3.3', 'API Endpoints'),
    ('3.4', 'ML Inference Pipeline'),
    ('3.5', 'PDF Report Generation'),
    ('4', 'Frontend Development'),
    ('4.1', 'Pages & Routes'),
    ('4.2', 'Authentication Modal'),
    ('4.3', 'Camera Analysis'),
    ('4.4', 'Result Page'),
    ('4.5', 'Dashboard'),
    ('5', 'Bugs Fixed'),
    ('6', 'Test Cases & Results'),
    ('6.1', 'Backend API Tests'),
    ('6.2', 'Auth Flow Tests'),
    ('6.3', 'Frontend Build Tests'),
    ('7', 'Known Limitations'),
    ('8', 'Deployment Instructions'),
]
for num, item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.3 * num.count('.') + 0.3)
    r = p.add_run(f'{num}   {item}')
    r.font.size = Pt(11)
    if '.' not in num:
        r.bold = True

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 1 — PROJECT OVERVIEW
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '1. Project Overview', level=1)
add_para(doc,
    'BattingEdge is a free, AI-powered cricket batting analysis web application developed '
    'as a Final Year Project at Bahria University, Karachi. It is Pakistan\'s first '
    'AI-based batting coaching platform, designed to democratise access to professional '
    'coaching insights for young Pakistani cricketers who cannot afford traditional coaching.'
)
doc.add_paragraph()

add_heading(doc, 'Problem Statement', level=2)
add_bullet(doc, 'Professional cricket coaching costs Rs. 3,000–10,000 per month in Pakistan')
add_bullet(doc, 'Less than 1% of aspiring cricketers have access to professional coaches')
add_bullet(doc, 'No free, AI-powered batting analysis tool exists for the Pakistani market')
add_bullet(doc, 'Young talent from rural and low-income backgrounds has no coaching pathway')

doc.add_paragraph()
add_heading(doc, 'Solution', level=2)
add_bullet(doc, 'Upload a batting video → AI analyses the shot type and biomechanics')
add_bullet(doc, 'Returns: shot classification, form score, grade, strengths, improvements, drills')
add_bullet(doc, 'ECB/MCC standard comparison for each body checkpoint')
add_bullet(doc, 'PDF report generation for offline review')
add_bullet(doc, 'Camera-based live recording with real-time MediaPipe pose detection')
add_bullet(doc, 'User authentication (email/password + Google OAuth)')
add_bullet(doc, 'Analysis history dashboard')

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 2 — ARCHITECTURE
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '2. System Architecture', level=1)

add_heading(doc, 'Technology Stack', level=2)
add_table(doc,
    ['Layer', 'Technology', 'Purpose'],
    [
        ['Backend Framework', 'FastAPI (Python 3.11)', 'REST API, async processing'],
        ['ML Model',          'TensorFlow / Keras BiLSTM', 'Shot classification'],
        ['ML Ensemble',       'XGBoost + Random Forest + Meta-Model', 'Stacking ensemble (94.71%)'],
        ['Pose Estimation',   'MediaPipe Pose (YOLO v8n)', 'Body keypoint extraction'],
        ['Database',          'SQLite (raw sqlite3)', 'Analyses + user data'],
        ['Auth',              'JWT (python-jose) + bcrypt', 'Secure token auth'],
        ['Email',             'Gmail SMTP SSL', 'Verification + reset emails'],
        ['Google OAuth',      'Google Identity Services (GSI)', 'One-tap Google Sign-In'],
        ['PDF Reports',       'ReportLab', 'Downloadable analysis reports'],
        ['Frontend',          'React 18 + Vite', 'SPA, fast HMR'],
        ['Styling',           'Tailwind CSS v4', 'Utility-first dark theme'],
        ['Animation',         'Framer Motion', 'Page + scroll animations'],
        ['Icons',             'Lucide React', 'Consistent icon set'],
        ['HTTP Client',       'Axios (with interceptors)', 'JWT auto-attach + 401 logout'],
        ['Notifications',     'react-hot-toast', 'In-app toast messages'],
    ]
)

doc.add_paragraph()
add_heading(doc, 'Data Flow', level=2)
steps = [
    '1. User uploads video (MP4/AVI/WebM etc.) via /api/upload',
    '2. Video saved to backend/uploads/ with UUID filename',
    '3. POST /api/analyze/{video_id} queues background task',
    '4. Background task: MediaPipe extracts 107 pose features × 50 frames',
    '5. Stacking ensemble classifies shot type (5 classes)',
    '6. Overlay video generated with skeleton annotations',
    '7. PDF report generated with ReportLab',
    '8. Results saved to SQLite analyses table',
    '9. Frontend polls GET /api/result/{video_id} every 2 seconds',
    '10. Result page displays: prediction, confidence, score, grade, biomechanical checks, drills',
]
for s in steps:
    add_bullet(doc, s)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 3 — BACKEND
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '3. Backend Development', level=1)
add_para(doc, 'The backend is built with FastAPI and serves as a REST API for the frontend. '
    'All ML inference is handled server-side. The backend uses raw sqlite3 (no ORM) as required '
    'by project constraints.')

add_heading(doc, '3.1 Database Schema', level=2)

add_heading(doc, 'analyses table', level=3)
add_table(doc,
    ['Column', 'Type', 'Description'],
    [
        ['video_id',         'TEXT PK', 'UUID v4 — unique per upload'],
        ['filename',         'TEXT',    'Original uploaded filename'],
        ['shot_type',        'TEXT',    'Classified shot (e.g. cover_drive)'],
        ['confidence',       'REAL',    'Model confidence percentage'],
        ['form_score',       'INTEGER', 'Overall form score /100'],
        ['all_probabilities','TEXT',    'JSON — per-class probabilities'],
        ['form_checks',      'TEXT',    'JSON — full form analysis object'],
        ['original_path',    'TEXT',    'Absolute path to uploaded video'],
        ['overlay_path',     'TEXT',    'Absolute path to annotated video'],
        ['status',           'TEXT',    'uploaded | processing | completed | failed'],
        ['error_message',    'TEXT',    'Error detail if failed'],
        ['progress',         'INTEGER', '0–100 processing progress'],
        ['created_at',       'TIMESTAMP','ISO datetime of upload'],
    ]
)

doc.add_paragraph()
add_heading(doc, 'users table', level=3)
add_table(doc,
    ['Column', 'Type', 'Description'],
    [
        ['user_id',       'TEXT PK', 'UUID v4'],
        ['email',         'TEXT UNIQUE', 'Normalised (lowercase, stripped)'],
        ['name',          'TEXT',   'Full name'],
        ['password_hash', 'TEXT',   'bcrypt hash (empty string for Google users)'],
        ['is_active',     'INTEGER','1 = active, 0 = deactivated'],
        ['is_verified',   'INTEGER','1 = email verified'],
        ['google_id',     'TEXT',   'Google sub (OAuth ID)'],
        ['created_at',    'TEXT',   'ISO datetime'],
    ]
)

doc.add_paragraph()
add_heading(doc, '3.2 Authentication System', level=2)
add_para(doc, 'Full JWT-based authentication with email verification and Google OAuth. '
    'Implemented in auth_models.py, auth_router.py, and auth_utils.py.')

add_table(doc,
    ['Feature', 'Implementation'],
    [
        ['Password hashing',    'bcrypt (direct — passlib incompatible with bcrypt 5.x)'],
        ['JWT tokens',          'python-jose, HS256, 7-day access token'],
        ['Verification tokens', 'JWT with type="verification", 24-hour expiry'],
        ['Reset tokens',        'JWT with type="reset", 1-hour expiry'],
        ['Email sending',       'Gmail SMTP SSL port 465, App Password'],
        ['Google OAuth',        'google-auth library, ID token verification'],
        ['Schema migration',    'PRAGMA table_info + ALTER TABLE (no drops ever)'],
    ]
)

doc.add_paragraph()
add_heading(doc, '3.3 API Endpoints', level=2)
add_table(doc,
    ['Method', 'Path', 'Auth', 'Description'],
    [
        ['GET',  '/',                         'No',  'API info and endpoint map'],
        ['GET',  '/api/health',               'No',  'Health check + model status'],
        ['POST', '/api/upload',               'No',  'Upload video file (multipart)'],
        ['POST', '/api/analyze/{video_id}',   'No',  'Start background analysis'],
        ['GET',  '/api/result/{video_id}',    'No',  'Poll analysis status/result'],
        ['GET',  '/api/video/{video_id}/overlay', 'No', 'Stream annotated video'],
        ['GET',  '/api/report/{video_id}/pdf','No',  'Download PDF report'],
        ['POST', '/auth/register',            'No',  'Register with email+password'],
        ['POST', '/auth/login',               'No',  'Login, returns JWT'],
        ['POST', '/auth/google',              'No',  'Google OAuth token exchange'],
        ['GET',  '/auth/verify-email',        'No',  'Verify email via token link'],
        ['GET',  '/auth/me',                  'JWT', 'Get current user profile'],
        ['PUT',  '/auth/update-password',     'JWT', 'Change password'],
        ['PUT',  '/auth/update-email',        'JWT', 'Change email'],
        ['POST', '/auth/forgot-password',     'No',  'Send password reset email'],
        ['POST', '/auth/reset-password',      'No',  'Reset password with token'],
        ['POST', '/auth/resend-verification', 'No',  'Resend verification email'],
    ]
)

doc.add_paragraph()
add_heading(doc, '3.4 ML Inference Pipeline (V9.5)', level=2)
add_para(doc, 'The inference pipeline is implemented in inference_v9_5.py (DO NOT MODIFY). '
    'It uses a stacking ensemble of four models:')
add_bullet(doc, 'BiLSTM (91.80% accuracy) — primary temporal sequence model')
add_bullet(doc, 'BiGRU (91.80% accuracy) — secondary temporal model')
add_bullet(doc, 'XGBoost (87.30% accuracy) — gradient boosted trees')
add_bullet(doc, 'Random Forest (89.42% accuracy) — ensemble trees')
add_bullet(doc, 'Meta-model (Logistic Regression) — stacks all four predictions → 94.71%')
doc.add_paragraph()
add_bullet(doc, 'Input: 50 video frames × 107 biomechanical features')
add_bullet(doc, 'Output: 5 shot classes (cover_drive, cut_shot, defense, pull_shot, sweep_shot)')
add_bullet(doc, 'Overlay: YOLO v8n person detection + MediaPipe skeleton drawn on video')

doc.add_paragraph()
add_heading(doc, '3.5 PDF Report Generation', level=2)
add_para(doc, 'Generated by report.py using ReportLab. Includes: shot classification, '
    'confidence score, form score, grade, ECB/MCC standards table, biomechanical breakdown, '
    'strengths, areas for improvement, and recommended training drills.')

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 4 — FRONTEND
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '4. Frontend Development', level=1)
add_para(doc, 'The frontend is a React 18 Single Page Application (SPA) built with Vite. '
    'Design theme: jet-black (#0a0a0f) background, neon-blue (#00d4ff) and neon-green (#00ff88) accents. '
    'Fonts: Space Grotesk (headings), Inter (body).')

add_heading(doc, '4.1 Pages & Routes', level=2)
add_table(doc,
    ['Route', 'Component', 'Description'],
    [
        ['/',                'LandingPage.jsx',       'Pakistan-focused hero, features, shot library, tech stats'],
        ['/analyze',         'AnalyzePage.jsx',       'Upload tab + Camera tab, 4-step progress stepper, polling'],
        ['/result/:videoId', 'ResultPage.jsx',        'Full analysis result with video, charts, drills'],
        ['/dashboard',       'DashboardPage.jsx',     'Analysis history from localStorage'],
        ['/verify',          'VerifyEmailPage.jsx',   'Email verification token handler'],
        ['/reset-password',  'ResetPasswordPage.jsx', 'Password reset form'],
    ]
)

doc.add_paragraph()
add_heading(doc, '4.2 Authentication Modal (AuthModal.jsx)', level=2)
add_bullet(doc, 'Login and Sign Up tabs, switchable')
add_bullet(doc, 'Google One-Tap Sign-In via GSI (accounts.google.com/gsi/client)')
add_bullet(doc, 'Forgot password flow → sends reset email → shows confirmation screen')
add_bullet(doc, 'Verify sent screen with resend link after registration')
add_bullet(doc, 'JWT stored in localStorage as be_token')
add_bullet(doc, 'AuthContext (React Context) provides user state across entire app')
add_bullet(doc, 'Axios interceptor auto-attaches Bearer token to all requests')
add_bullet(doc, 'Auto-logout on 401 — dispatches be:logout event')

doc.add_paragraph()
add_heading(doc, '4.3 Camera Analysis (CameraModal.jsx)', level=2)
add_bullet(doc, 'getUserMedia API — 720p camera stream, mirrored for natural view')
add_bullet(doc, 'MediaPipe Pose loaded from CDN (legacy @mediapipe/pose@0.5)')
add_bullet(doc, '33-point pose skeleton drawn with raw Canvas API (neon-blue lines, neon-green joints)')
add_bullet(doc, '"Loading AI Pose…" spinner badge while WASM model initialises (30–60s)')
add_bullet(doc, 'MediaRecorder API — records in WebM/VP8, max 60 seconds')
add_bullet(doc, 'Auto-stop at 60s, retake option before submitting')
add_bullet(doc, 'Recorded blob → File object → uploaded via /api/upload')

doc.add_paragraph()
add_heading(doc, '4.4 Result Page (ResultPage.jsx)', level=2)
add_bullet(doc, 'Real-time polling every 2 seconds against /api/result/{videoId}')
add_bullet(doc, 'Animated shot header: detected shot, confidence, score /100, grade badge, level')
add_bullet(doc, 'Low-confidence warning banner (< 50%)')
add_bullet(doc, 'Shot Probability Breakdown — animated bars for all 5 shot classes')
add_bullet(doc, 'Overlay video player (annotated skeleton video)')
add_bullet(doc, 'Strengths and To-Improve two-column grid')
add_bullet(doc, 'Coach\'s Assessment (AI-generated summary paragraph)')
add_bullet(doc, 'Biomechanical Breakdown — expandable check items with measured values vs ECB/MCC targets')
add_bullet(doc, 'Recommended Drills — 3 drills with icons and descriptions')
add_bullet(doc, 'Shot Facts panel — 3 Pakistani players to study + coaching tip + fun fact per shot')
add_bullet(doc, 'PDF download banner')

doc.add_paragraph()
add_heading(doc, '4.5 Dashboard (DashboardPage.jsx)', level=2)
add_bullet(doc, 'Reads analysis history from localStorage (be_history key, max 20 entries)')
add_bullet(doc, 'Stats cards: total analyses, average score, best grade, latest date')
add_bullet(doc, 'History table: shot type, score, grade, date, View + PDF links')

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 5 — BUGS FIXED
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '5. Bugs Fixed', level=1)
add_table(doc,
    ['#', 'Bug', 'Root Cause', 'Fix Applied'],
    [
        ['1', 'CORS blocking port 5174',
              'Dev server moved from 5173 to 5174; old allow_origins list',
              'Added 5174, 5175, 5176 to CORSMiddleware allow_origins in main.py'],
        ['2', 'Google OAuth field mismatch',
              'Frontend sent {token} but backend expected {google_token}',
              'Fixed authGoogle() in api.js to send {google_token: token}'],
        ['3', 'password_hash NOT NULL breaks Google users',
              'database.py created users table with NOT NULL on password_hash before auth_models',
              'Removed CREATE TABLE users from database.py; auth_models now owns it with nullable column'],
        ['4', 'Google users crash on null password_hash',
              'create_user() passed None → sqlite3 NOT NULL constraint error',
              'safe_pw_hash = password_hash if password_hash is not None else "" in auth_models.py'],
        ['5', 'ResultPage showed blank data',
              'Old code read data.form_checks but API wraps result under data.result',
              'Fixed data path to data.result.form_analysis throughout ResultPage'],
        ['6', 'Probability bars showed 8365%, 616% etc.',
              'Inference returns 0–100 range; code multiplied by 100 again',
              'pct = v > 1 ? Math.round(v) : Math.round(v * 100)'],
        ['7', '.env not loading → GOOGLE_CLIENT_ID missing',
              'load_dotenv() searched CWD (project root) but .env is in backend/',
              'load_dotenv(Path(__file__).parent / ".env") in both main.py and auth_utils.py'],
        ['8', 'Verification email link shows blank page',
              '/verify route did not exist in React router',
              'Created VerifyEmailPage.jsx and added /verify route to App.jsx'],
        ['9', '/reset-password route missing',
              'Reset email linked to frontend page that did not exist',
              'Created ResetPasswordPage.jsx and added /reset-password route'],
        ['10', 'MediaPipe skeleton not visible',
              'Relied on window.drawConnectors from CDN — not always exported as window global',
              'Replaced with raw Canvas API + hardcoded 33-joint CONNECTIONS array'],
        ['11', 'Port 8000 already in use on restart',
              'Previous server instance still running on same port',
              'Kill via PowerShell Stop-Process before restarting'],
    ]
)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 6 — TEST CASES
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '6. Test Cases & Results', level=1)

add_heading(doc, '6.1 Backend API Tests', level=2)
add_para(doc, 'All tests executed via curl against running server on localhost:8000.')

add_table(doc,
    ['TC#', 'Endpoint', 'Input', 'Expected', 'Result'],
    [
        ['TC01', 'GET /api/health',          '—',                           '200, model_loaded=true, accuracy=94.71%',   'PASS'],
        ['TC02', 'GET /',                     '—',                           '200, endpoint map returned',                'PASS'],
        ['TC03', 'POST /api/upload',          'Valid .mp4 file',             '200, video_id UUID returned',               'PASS'],
        ['TC04', 'POST /api/upload',          '.txt file (wrong type)',      '400, Invalid file type error',              'PASS'],
        ['TC05', 'POST /api/upload',          'File > 100MB',                '413, File too large error',                 'PASS'],
        ['TC06', 'GET /api/result/:id',       'Valid video_id (uploaded)',   '200, status=uploaded, progress=0',          'PASS'],
        ['TC07', 'GET /api/result/:id',       'Unknown video_id',            '404, Analysis not found',                   'PASS'],
        ['TC08', 'GET /api/report/:id/pdf',   'Non-existent PDF',            '404, PDF not found',                        'PASS'],
        ['TC09', 'GET /api/video/:id/overlay','No overlay generated yet',    '404, overlay not ready',                    'PASS'],
    ]
)

doc.add_paragraph()
add_heading(doc, '6.2 Auth Flow Tests', level=2)
add_table(doc,
    ['TC#', 'Test', 'Input', 'Expected', 'Result'],
    [
        ['TC10', 'Register new user',          'Valid email, password ≥ 8 chars', '201, registration successful',              'PASS'],
        ['TC11', 'Register duplicate email',   'Already registered email',        '400, Email already registered',             'PASS'],
        ['TC12', 'Register short password',    'Password < 8 chars',              '400, Password too short',                   'PASS'],
        ['TC13', 'Register invalid email',     'email without @ or .',            '400, Invalid email format',                 'PASS'],
        ['TC14', 'Login before verification',  'Valid credentials, unverified',   '403, Please verify your email',             'PASS'],
        ['TC15', 'Login after verification',   'Valid credentials, verified',     '200, access_token returned',                'PASS'],
        ['TC16', 'Login wrong password',       'Wrong password',                  '401, Invalid email or password',            'PASS'],
        ['TC17', 'Login unknown email',        'Non-existent email',              '401, Invalid email or password',            'PASS'],
        ['TC18', 'GET /auth/me valid token',   'Valid Bearer JWT',                '200, user profile with is_verified=true',   'PASS'],
        ['TC19', 'GET /auth/me invalid token', 'Malformed JWT',                   '401, Invalid or expired token',             'PASS'],
        ['TC20', 'GET /auth/me no token',      'No Authorization header',         '401, Not authenticated',                    'PASS'],
        ['TC21', 'Update password',            'Valid current + new password',    '200, Password updated successfully',        'PASS'],
        ['TC22', 'Login with new password',    'New password after TC21',         '200, access_token returned',                'PASS'],
        ['TC23', 'Forgot password',            'Registered email',                '200, safe message (no user enumeration)',    'PASS'],
        ['TC24', 'Forgot password unknown',    'Unregistered email',              '200, same safe message',                    'PASS'],
        ['TC25', 'Resend verification',        'Any email',                       '200, safe message',                         'PASS'],
        ['TC26', 'Email verification link',    '/verify?token=valid_token',       'Frontend VerifyEmailPage renders, API called','PASS'],
        ['TC27', 'Expired verify token',       'Expired JWT token',               '400, Invalid or expired verification link', 'PASS'],
        ['TC28', 'CORS headers port 5173',     'Origin: localhost:5173',          'access-control-allow-origin: localhost:5173','PASS'],
        ['TC29', 'CORS headers port 5174',     'Origin: localhost:5174',          'access-control-allow-origin: localhost:5174','PASS'],
        ['TC30', 'Google OAuth wrong token',   'Invalid google_token',            '400, Invalid Google token',                 'PASS'],
        ['TC31', 'Google OAuth missing field', '{token} instead of {google_token}','422, validation error',                   'PASS'],
    ]
)

doc.add_paragraph()
add_heading(doc, '6.3 Frontend Build Tests', level=2)
add_table(doc,
    ['TC#', 'Test', 'Result', 'Notes'],
    [
        ['TC32', 'npm run build — initial build',        'PASS', '2109 modules, 0 errors'],
        ['TC33', 'npm run build — after CORS fix',       'PASS', '2109 modules, 0 errors'],
        ['TC34', 'npm run build — after ResultPage fix', 'PASS', '0 errors'],
        ['TC35', 'npm run build — after CameraModal fix','PASS', '0 errors'],
        ['TC36', 'npm run build — after VerifyEmail add','PASS', '2111 modules, 0 errors'],
        ['TC37', 'npm run build — final state',          'PASS', '2111 modules, 433 KB JS (gzip 139 KB)'],
    ]
)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 7 — KNOWN LIMITATIONS
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '7. Known Limitations', level=1)
add_table(doc,
    ['Limitation', 'Detail', 'Mitigation'],
    [
        ['Overlay video shows ??? for angles',
         'OpenCV putText() does not support Unicode (degree symbol °)',
         'Known — inference_v9_5.py is off-limits; cosmetic issue only'],
        ['Model loads in ~25 seconds',
         'TensorFlow/Keras BiLSTM + XGBoost + RF load on server startup',
         'Server starts once; no cold load per request'],
        ['Classifies any video',
         'Model always outputs one of 5 classes even for non-cricket content',
         'Low-confidence warning shown in UI when confidence < 50%'],
        ['Google OAuth origin_mismatch',
         'Google Cloud Console requires authorized origins per domain/port',
         'User must add localhost:5173 and localhost:5174 in GCP console'],
        ['Gmail App Password numeric',
         'The stored GMAIL_APP_PASSWORD was all digits (invalid format)',
         'User must generate real 16-letter App Password from Google account'],
        ['MediaPipe WASM download time',
         'MediaPipe loads from CDN and takes 30–60 seconds on first open',
         '"Loading AI Pose…" spinner badge shown during load'],
        ['Analysis history is local-only',
         'History stored in localStorage — not synced with user account',
         'Post-defense: migrate to server-side history table'],
        ['sklearn version mismatch warning',
         'StandardScaler pickled with sklearn 1.7.2, server runs 1.8.0',
         'Warning only, not an error; results remain valid'],
    ]
)

page_break(doc)

# ─────────────────────────────────────────────────────────────────────────────
# SECTION 8 — DEPLOYMENT
# ─────────────────────────────────────────────────────────────────────────────

add_heading(doc, '8. Deployment Instructions', level=1)

add_heading(doc, 'Prerequisites', level=2)
add_bullet(doc, 'Python 3.11 with virtualenv')
add_bullet(doc, 'Node.js 18+ and npm')
add_bullet(doc, 'Google Cloud Console project with OAuth 2.0 Client ID')
add_bullet(doc, 'Gmail account with 2-Step Verification enabled')

add_heading(doc, 'Backend Setup', level=2)
add_code(doc, 'cd BattingEdge_FYP')
add_code(doc, 'python -m venv venv')
add_code(doc, 'venv\\Scripts\\activate')
add_code(doc, 'pip install -r requirements.txt')
add_code(doc, '# Edit backend/.env with your GMAIL_APP_PASSWORD and GOOGLE_CLIENT_ID')
add_code(doc, 'cd backend && python main.py')

add_heading(doc, 'Frontend Setup', level=2)
add_code(doc, 'cd frontend')
add_code(doc, 'npm install')
add_code(doc, 'npm run dev')

add_heading(doc, 'Google Auth Setup', level=2)
items = [
    '1. Go to console.cloud.google.com → APIs & Services → Credentials',
    '2. Open your OAuth 2.0 Client ID',
    '3. Add to Authorized JavaScript origins: http://localhost:5173 and http://localhost:5174',
    '4. Add to Authorized redirect URIs: http://localhost:5173 and http://localhost:5174',
    '5. Save — propagation takes up to 5 minutes',
]
for item in items:
    add_bullet(doc, item)

add_heading(doc, 'Gmail App Password Setup', level=2)
items2 = [
    '1. Go to myaccount.google.com/security → Enable 2-Step Verification',
    '2. Go to myaccount.google.com/apppasswords',
    '3. Create an App Password named "BattingEdge"',
    '4. Copy the 16-letter password (format: xxxx xxxx xxxx xxxx)',
    '5. Paste into backend/.env as: GMAIL_APP_PASSWORD=xxxxxxxxxxxxxxxx (no spaces)',
]
for item in items2:
    add_bullet(doc, item)

add_heading(doc, 'backend/.env Template', level=2)
add_code(doc, 'SECRET_KEY=<generate with: python -c "import secrets; print(secrets.token_hex(32))">')
add_code(doc, 'ALGORITHM=HS256')
add_code(doc, 'GMAIL_ADDRESS=your@gmail.com')
add_code(doc, 'GMAIL_APP_PASSWORD=abcdefghijklmnop  # 16 letters, no spaces')
add_code(doc, 'GOOGLE_CLIENT_ID=your-client-id.apps.googleusercontent.com')
add_code(doc, 'FRONTEND_URL=http://localhost:5173')

# ─────────────────────────────────────────────────────────────────────────────
# SAVE
# ─────────────────────────────────────────────────────────────────────────────

out = r'D:\Users\Anoshia\BattingEdge_FYP\BattingEdge_Project_Report.docx'
doc.save(out)
print(f'Report saved: {out}')
